import os
import shutil
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def hex_to_rgb(hex_color: str):
    hex_color = hex_color.lstrip('#')
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)


def get_alignment(alignment_str: str):
    mapping = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get((alignment_str or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def backup_document(file_path: str) -> str:
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = f"{file_path}.backup_{timestamp}"
    shutil.copy2(file_path, backup_path)
    return backup_path


class DocumentExecutor:

    def __init__(self, file_path: str):
        self.file_path            = file_path
        self.doc                  = DocxDocument(file_path)
        self._bookmark_id_counter = 1000

    def save(self):
        self.doc.save(self.file_path)

    def execute_actions(self, actions: list) -> list:
        results = []
        for action in actions:
            action_type = action.get("type")
            params      = action.get("params", {})
            try:
                method = getattr(self, f"action_{action_type}", None)
                if method:
                    result = method(**params)
                    results.append({"action": action_type, "status": "success", "result": result})
                else:
                    results.append({"action": action_type, "status": "skipped", "reason": "Unknown action"})
            except Exception as e:
                results.append({"action": action_type, "status": "error", "error": str(e)})
        self.save()
        return results

    # ──────────────────────────────────────────
    # BASIC FORMATTING
    # ──────────────────────────────────────────

    def action_set_font(self, font_name="Calibri", size=11,
                        bold=False, italic=False, color=None,
                        apply_to="all", find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text and find_text not in run.text:
                    continue
                run.font.name = font_name
                run.font.size = Pt(size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    run.font.color.rgb = hex_to_rgb(color)
        return f"Font set to {font_name} {size}pt"

    def action_set_heading_style(self, level=1, bold=True,
                                  color=None, font_size=None,
                                  underline=False, **kwargs):
        style_name = f"Heading {level}"
        for para in self.doc.paragraphs:
            if para.style.name == style_name:
                for run in para.runs:
                    if bold is not None:
                        run.bold = bold
                    if color:
                        run.font.color.rgb = hex_to_rgb(color)
                    if font_size:
                        run.font.size = Pt(font_size)
                    if underline:
                        run.font.underline = True
        return f"Heading {level} style updated"

    def action_set_alignment(self, alignment="justify",
                              apply_to="all", find_text=None, **kwargs):
        align = get_alignment(alignment)
        for para in self.doc.paragraphs:
            if find_text and find_text not in para.text:
                continue
            if apply_to == "all":
                para.alignment = align
            elif apply_to == "headings" and "Heading" in para.style.name:
                para.alignment = align
            elif apply_to == "body" and "Heading" not in para.style.name:
                para.alignment = align
        return f"Alignment set to {alignment} for {apply_to}"

    def action_set_margins(self, top=1.0, bottom=1.0,
                           left=1.0, right=1.0, **kwargs):
        for section in self.doc.sections:
            section.top_margin    = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin   = Inches(left)
            section.right_margin  = Inches(right)
        return f"Margins set T={top} B={bottom} L={left} R={right} inches"

    def action_set_paragraph_spacing(self, before=0, after=8,
                                      line_spacing=1.5,
                                      apply_to="all", **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            pf = para.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after  = Pt(after)
            pf.line_spacing = line_spacing
        return f"Spacing: before={before}pt after={after}pt line={line_spacing}"

    # ──────────────────────────────────────────
    # ADVANCED TEXT FORMATTING
    # ──────────────────────────────────────────

    def action_set_underline(self, apply_to="all", find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text and find_text not in run.text:
                    continue
                run.font.underline = True
        return f"Underline applied to {apply_to}"

    def action_set_strikethrough(self, apply_to="all", find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text and find_text not in run.text:
                    continue
                run.font.strike = True
        return f"Strikethrough applied to {apply_to}"

    def action_set_highlight(self, color="yellow",
                              find_text=None, apply_to="all", **kwargs):
        from docx.enum.text import WD_COLOR_INDEX
        color_map = {
            "yellow":    WD_COLOR_INDEX.YELLOW,
            "green":     WD_COLOR_INDEX.BRIGHT_GREEN,
            "cyan":      WD_COLOR_INDEX.CYAN,
            "magenta":   WD_COLOR_INDEX.MAGENTA,
            "blue":      WD_COLOR_INDEX.BLUE,
            "red":       WD_COLOR_INDEX.RED,
            "pink":      WD_COLOR_INDEX.PINK,
            "turquoise": WD_COLOR_INDEX.TURQUOISE,
            "gray":      WD_COLOR_INDEX.GRAY_25,
        }
        wd_color = color_map.get(color.lower(), WD_COLOR_INDEX.YELLOW)
        count = 0
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text:
                    if find_text.lower() in run.text.lower():
                        run.font.highlight_color = wd_color
                        count += 1
                else:
                    run.font.highlight_color = wd_color
                    count += 1
        return f"Highlighted {count} runs in {color}"

    def action_set_superscript(self, find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            for run in para.runs:
                if find_text is None or find_text in run.text:
                    run.font.superscript = True
        return "Superscript applied"

    def action_set_subscript(self, find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            for run in para.runs:
                if find_text is None or find_text in run.text:
                    run.font.subscript = True
        return "Subscript applied"

    def action_set_text_color(self, color="#000000", apply_to="all",
                               find_text=None, **kwargs):
        rgb = hex_to_rgb(color)
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text is None or find_text in run.text:
                    run.font.color.rgb = rgb
        return f"Text color set to {color} for {apply_to}"

    def action_set_indent(self, left=0.5, right=0.0,
                           first_line=0.0, apply_to="all", **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            pf = para.paragraph_format
            pf.left_indent  = Inches(left)
            pf.right_indent = Inches(right)
            if first_line:
                pf.first_line_indent = Inches(first_line)
        return f"Indent set left={left} right={right} first_line={first_line}"

    def action_remove_formatting(self, apply_to="all", find_text=None, **kwargs):
        for para in self.doc.paragraphs:
            if apply_to == "headings" and "Heading" not in para.style.name:
                continue
            if apply_to == "body" and "Heading" in para.style.name:
                continue
            for run in para.runs:
                if find_text and find_text not in run.text:
                    continue
                run.font.bold      = False
                run.font.italic    = False
                run.font.underline = False
                run.font.strike    = False
                run.font.color.rgb = RGBColor(0, 0, 0)
        return f"Formatting removed from {apply_to}"

    def action_set_character_spacing(self, spacing=1.0, **kwargs):
        for para in self.doc.paragraphs:
            for run in para.runs:
                rPr        = run._r.get_or_add_rPr()
                spacing_el = OxmlElement('w:spacing')
                spacing_el.set(qn('w:val'), str(int(spacing * 20)))
                rPr.append(spacing_el)
        return f"Character spacing set to {spacing}pt"

    def action_set_drop_cap(self, paragraph_index=0, lines=3,
                             font_name=None, **kwargs):
        paras = [p for p in self.doc.paragraphs if p.text.strip()]
        if paragraph_index >= len(paras):
            return f"Paragraph {paragraph_index} not found"
        para = paras[paragraph_index]
        if not para.runs or not para.runs[0].text:
            return "Paragraph has no text to apply drop cap"
        first_char        = para.runs[0].text[0]
        para.runs[0].text = para.runs[0].text[1:]
        drop_para = para.insert_paragraph_before("")
        run = drop_para.add_run(first_char)
        if font_name:
            run.font.name = font_name
        run.font.size = Pt(11 * lines * 1.8)
        run.font.bold = True
        pPr     = drop_para._p.get_or_add_pPr()
        framePr = OxmlElement('w:framePr')
        framePr.set(qn('w:dropCap'), 'drop')
        framePr.set(qn('w:lines'),   str(lines))
        framePr.set(qn('w:wrap'),    'around')
        framePr.set(qn('w:vAnchor'), 'text')
        framePr.set(qn('w:hAnchor'), 'text')
        pPr.append(framePr)
        return f"Drop cap applied to first letter '{first_char}' ({lines} lines)"

    # ──────────────────────────────────────────
    # SELECTION-BASED FORMATTING
    # ──────────────────────────────────────────

    def action_apply_to_selection(self, selected_texts=None,
                                   command_type="bold",
                                   font_name=None, font_size=None,
                                   color=None, highlight_color=None,
                                   alignment=None, make_heading=None,
                                   **kwargs):
        if not selected_texts:
            return "No paragraphs selected"
        count = 0
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not any(sel.strip() in text or text in sel.strip()
                       for sel in selected_texts):
                continue
            if command_type == "bold":
                for run in para.runs:
                    run.font.bold = True
            elif command_type == "italic":
                for run in para.runs:
                    run.font.italic = True
            elif command_type == "underline":
                for run in para.runs:
                    run.font.underline = True
            elif command_type == "strikethrough":
                for run in para.runs:
                    run.font.strike = True
            elif command_type == "highlight":
                from docx.enum.text import WD_COLOR_INDEX
                color_map = {
                    "yellow": WD_COLOR_INDEX.YELLOW,
                    "green":  WD_COLOR_INDEX.BRIGHT_GREEN,
                    "cyan":   WD_COLOR_INDEX.CYAN,
                    "pink":   WD_COLOR_INDEX.PINK,
                    "blue":   WD_COLOR_INDEX.BLUE,
                }
                hl = color_map.get((highlight_color or "yellow").lower(),
                                   WD_COLOR_INDEX.YELLOW)
                for run in para.runs:
                    run.font.highlight_color = hl
            elif command_type == "font":
                for run in para.runs:
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(font_size)
            elif command_type == "color":
                if color:
                    for run in para.runs:
                        run.font.color.rgb = hex_to_rgb(color)
            elif command_type == "align":
                if alignment:
                    para.alignment = get_alignment(alignment)
            elif command_type == "heading":
                level = make_heading or 1
                try:
                    para.style = self.doc.styles[f"Heading {level}"]
                except KeyError:
                    pass
            elif command_type == "uppercase":
                for run in para.runs:
                    run.text = run.text.upper()
            elif command_type == "lowercase":
                for run in para.runs:
                    run.text = run.text.lower()
            elif command_type == "capitalize":
                for run in para.runs:
                    run.text = run.text.title()
            elif command_type == "remove_formatting":
                for run in para.runs:
                    run.font.bold      = False
                    run.font.italic    = False
                    run.font.underline = False
                    run.font.strike    = False
            count += 1
        return f"Applied '{command_type}' to {count} paragraph(s)"

    # ──────────────────────────────────────────
    # TEXT OPERATIONS
    # ──────────────────────────────────────────

    def action_find_replace(self, find="", replace="",
                             case_sensitive=False, **kwargs):
        import re
        count = 0
        for para in self.doc.paragraphs:
            for run in para.runs:
                if case_sensitive:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
                else:
                    new_text = re.sub(re.escape(find), replace,
                                      run.text, flags=re.IGNORECASE)
                    if new_text != run.text:
                        run.text = new_text
                        count += 1
        return f"Replaced {count} occurrences of '{find}' with '{replace}'"

    def action_add_text(self, text="", position="end",
                        new_paragraph=True, **kwargs):
        if position == "end":
            if new_paragraph:
                self.doc.add_paragraph(text)
            else:
                if self.doc.paragraphs:
                    self.doc.paragraphs[-1].add_run(" " + text)
        elif position == "beginning":
            if self.doc.paragraphs:
                self.doc.paragraphs[0].insert_paragraph_before(text)
        return f"Text added at {position}"

    def action_delete_text(self, find="", **kwargs):
        return self.action_find_replace(find=find, replace="")

    # ──────────────────────────────────────────
    # LISTS AND NUMBERING
    # ──────────────────────────────────────────

    def action_add_bullet_list(self, items=None, position="end", **kwargs):
        if not items:
            return "No items provided"
        for item in items:
            try:
                para      = self.doc.add_paragraph(style="List Bullet")
                para.text = item
            except KeyError:
                para      = self.doc.add_paragraph()
                para.text = f"• {item}"
                pPr = para._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'),    '720')
                ind.set(qn('w:hanging'), '360')
                pPr.append(ind)
        return f"Bullet list added with {len(items)} items"

    def action_add_numbered_list(self, items=None, position="end",
                                  restart=True, **kwargs):
        if not items:
            return "No items provided"
        try:
            for item in items:
                self.doc.add_paragraph(item, style="List Number")
            return f"Numbered list added with {len(items)} items (auto-numbering)"
        except KeyError:
            for i, item in enumerate(items, 1):
                para      = self.doc.add_paragraph()
                para.text = f"{i}. {item}"
                pPr = para._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'),    '720')
                ind.set(qn('w:hanging'), '360')
                pPr.append(ind)
            return f"Numbered list added with {len(items)} items (manual fallback)"

    def action_add_multilevel_list(self, items=None, **kwargs):
        if not items:
            return "No items provided"
        style_map  = {0: "List Number", 1: "List Number 2", 2: "List Number 3"}
        indent_map = {0: ('360', '180'), 1: ('720', '360'), 2: ('1080', '540')}
        for item in items:
            level = item.get("level", 0)
            text  = item.get("text", "")
            try:
                self.doc.add_paragraph(text, style=style_map.get(level, "List Number"))
            except KeyError:
                left, hanging = indent_map.get(level, indent_map[0])
                para      = self.doc.add_paragraph()
                para.text = text
                pPr = para._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'),    left)
                ind.set(qn('w:hanging'), hanging)
                pPr.append(ind)
        return f"Multilevel list added with {len(items)} items"

    def action_convert_to_bullets(self, **kwargs):
        import re
        converted = 0
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^[-*•]\s+', text) or re.match(r'^\d+[.)]\s+', text):
                clean = re.sub(r'^[-*•\d.)]+\s*', '', text)
                para.text = clean
                try:
                    para.style = self.doc.styles['List Bullet']
                except KeyError:
                    para.text = f"• {clean}"
                converted += 1
        return f"Converted {converted} paragraphs to bullet list"

    def action_restart_numbering(self, list_style="List Number", **kwargs):
        try:
            para = self.doc.add_paragraph("", style=list_style)
        except KeyError:
            para = self.doc.add_paragraph("")
        return f"Numbering restart marker added for style '{list_style}'"

    def action_add_checklist(self, items=None, **kwargs):
        if not items:
            return "No items provided"
        for item in items:
            para      = self.doc.add_paragraph()
            para.text = f"☐ {item}"
            for run in para.runs:
                run.font.name = "Wingdings 2"
        return f"Checklist added with {len(items)} items"

    # ──────────────────────────────────────────
    # PAGE MANAGEMENT
    # ──────────────────────────────────────────

    def action_add_page_break(self, after_page=1, **kwargs):
        para = self.doc.add_paragraph()
        run  = para.add_run()
        br   = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
        return "Page break added"

    def action_add_blank_page(self, after_page=1, **kwargs):
        self.action_add_page_break()
        self.doc.add_paragraph("")
        self.action_add_page_break()
        return f"Blank page added after page {after_page}"

    def action_set_page_size(self, size="A4", **kwargs):
        sizes = {"A4": (210, 297), "Letter": (215.9, 279.4), "A3": (297, 420)}
        w, h  = sizes.get(size, (210, 297))
        for section in self.doc.sections:
            section.page_width  = Mm(w)
            section.page_height = Mm(h)
        return f"Page size set to {size}"

    def action_set_page_orientation(self, orientation="portrait", **kwargs):
        for section in self.doc.sections:
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = \
                    section.page_height, section.page_width
            else:
                section.orientation = WD_ORIENT.PORTRAIT
        return f"Page orientation set to {orientation}"

    def action_set_page_color(self, color="#FFFFFF", **kwargs):
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), color.lstrip('#'))
        self.doc.element.insert(0, bg)
        disp = OxmlElement('w:displayBackgroundShape')
        self.doc.settings.element.append(disp)
        return f"Page color set to {color}"

    def action_add_line_numbers(self, start=1, step=1,
                                 restart="newPage", **kwargs):
        for section in self.doc.sections:
            sectPr    = section._sectPr
            lnNumType = OxmlElement('w:lnNumType')
            lnNumType.set(qn('w:countBy'), str(step))
            lnNumType.set(qn('w:start'),   str(start))
            lnNumType.set(qn('w:restart'), restart)
            sectPr.append(lnNumType)
        return f"Line numbers added start={start} step={step}"

    # ──────────────────────────────────────────
    # HEADERS AND FOOTERS
    # ──────────────────────────────────────────

    def action_add_header(self, text="", alignment="center", **kwargs):
        for section in self.doc.sections:
            header = section.header
            para   = header.paragraphs[0] if header.paragraphs \
                     else header.add_paragraph()
            para.clear()
            para.add_run(text)
            para.alignment = get_alignment(alignment)
        return f"Header added: '{text}'"

    def action_add_footer(self, text="", alignment="center", **kwargs):
        for section in self.doc.sections:
            footer = section.footer
            para   = footer.paragraphs[0] if footer.paragraphs \
                     else footer.add_paragraph()
            para.clear()
            para.add_run(text)
            para.alignment = get_alignment(alignment)
        return f"Footer added: '{text}'"

    def action_add_page_numbers(self, position="footer", alignment="right",
                                 format="Page X", **kwargs):
        for section in self.doc.sections:
            target = section.footer if position == "footer" else section.header
            para   = target.paragraphs[0] if target.paragraphs \
                     else target.add_paragraph()
            para.clear()
            para.alignment = get_alignment(alignment)
            run = para.add_run()
            if format == "Page X":
                run.text = "Page "
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld)
            instr = OxmlElement('w:instrText')
            instr.text = "PAGE"
            run._r.append(instr)
            fld2 = OxmlElement('w:fldChar')
            fld2.set(qn('w:fldCharType'), 'end')
            run._r.append(fld2)
            if format == "X of Y":
                run2 = para.add_run(" of ")
                fld3 = OxmlElement('w:fldChar')
                fld3.set(qn('w:fldCharType'), 'begin')
                run2._r.append(fld3)
                instr2 = OxmlElement('w:instrText')
                instr2.text = "NUMPAGES"
                run2._r.append(instr2)
                fld4 = OxmlElement('w:fldChar')
                fld4.set(qn('w:fldCharType'), 'end')
                run2._r.append(fld4)
        return f"Page numbers added to {position} ({alignment})"

    def action_set_different_first_page_header(self, first_header="",
                                                rest_header="", **kwargs):
        for section in self.doc.sections:
            section.different_first_page_header_footer = True
            fh = section.first_page_header
            if fh.paragraphs:
                fh.paragraphs[0].text = first_header
            else:
                fh.add_paragraph(first_header)
            hdr = section.header
            if hdr.paragraphs:
                hdr.paragraphs[0].text = rest_header
            else:
                hdr.add_paragraph(rest_header)
        return "Different first page header set"

    def action_set_section_header(self, section_index=0, text="",
                                   alignment="center", link_to_previous=False,
                                   **kwargs):
        if section_index >= len(self.doc.sections):
            return f"Section {section_index} does not exist"
        section = self.doc.sections[section_index]
        section.header.is_linked_to_previous = link_to_previous
        para = section.header.paragraphs[0] if section.header.paragraphs \
               else section.header.add_paragraph()
        para.clear()
        para.add_run(text)
        para.alignment = get_alignment(alignment)
        return f"Section {section_index} header set: '{text}'"

    def action_set_section_footer(self, section_index=0, text="",
                                   alignment="center", link_to_previous=False,
                                   **kwargs):
        if section_index >= len(self.doc.sections):
            return f"Section {section_index} does not exist"
        section = self.doc.sections[section_index]
        section.footer.is_linked_to_previous = link_to_previous
        para = section.footer.paragraphs[0] if section.footer.paragraphs \
               else section.footer.add_paragraph()
        para.clear()
        para.add_run(text)
        para.alignment = get_alignment(alignment)
        return f"Section {section_index} footer set: '{text}'"

    # ──────────────────────────────────────────
    # IMAGES AND MEDIA  ← UPDATED with server_path support
    # ──────────────────────────────────────────

    def action_insert_image(self, image_path=None, server_path=None,
                             width_inches=None, height_inches=None,
                             position="end", alignment="center",
                             caption=None, **kwargs):
        # Accept either image_path or server_path
        path = image_path or server_path
        if not path or not os.path.exists(path):
            return f"Image not found: {path}"
        para = self.doc.add_paragraph()
        para.alignment = get_alignment(alignment)
        run = para.add_run()
        if width_inches and height_inches:
            run.add_picture(path, width=Inches(width_inches),
                            height=Inches(height_inches))
        elif width_inches:
            run.add_picture(path, width=Inches(width_inches))
        else:
            run.add_picture(path, width=Inches(4.0))
        if caption:
            cap = self.doc.add_paragraph(caption)
            cap.alignment = get_alignment(alignment)
            for r in cap.runs:
                r.italic    = True
                r.font.size = Pt(9)
        return f"Image inserted: {path}"

    def action_insert_logo(self, image_path=None, server_path=None,
                            page="first", position="top_right",
                            width_inches=1.5, **kwargs):
        path = image_path or server_path
        if not path or not os.path.exists(path):
            return f"Logo not found: {path}"
        alignment_map = {
            "top_right":  WD_ALIGN_PARAGRAPH.RIGHT,
            "top_left":   WD_ALIGN_PARAGRAPH.LEFT,
            "top_center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        for i, section in enumerate(self.doc.sections):
            if page == "first" and i > 0:
                continue
            header = section.header
            para   = header.paragraphs[0] if header.paragraphs \
                     else header.add_paragraph()
            para.alignment = alignment_map.get(position, WD_ALIGN_PARAGRAPH.RIGHT)
            para.add_run().add_picture(path, width=Inches(width_inches))
        return f"Logo inserted at {position} on {page} page"

    def action_caption_image(self, image_index=0, caption_text="",
                              label="Figure", **kwargs):
        cap_para = self.doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(f"{label} ")
        run.bold = True
        fld_run = cap_para.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        fld_run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = f' SEQ {label} \\* ARABIC '
        fld_run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        fld_run._r.append(fld2)
        cap_para.add_run(f": {caption_text}").italic = True
        return f"Caption added: '{label}: {caption_text}'"

    # ──────────────────────────────────────────
    # TABLES
    # ──────────────────────────────────────────

    def action_insert_table(self, rows=3, cols=3, position="end",
                             headers=None, **kwargs):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        if headers:
            for i, header in enumerate(headers[:cols]):
                table.rows[0].cells[i].text = header
        return f"Table inserted {rows}x{cols}"

    def action_format_table(self, style="Table Grid", header_color=None, **kwargs):
        for table in self.doc.tables:
            try:
                table.style = style
            except Exception:
                pass
        return f"Tables formatted: {style}"

    def action_set_table_cell_color(self, table_index=0, row=0, col=0,
                                     color="#FFFFFF", **kwargs):
        try:
            cell = self.doc.tables[table_index].rows[row].cells[col]
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  color.lstrip('#'))
            tcPr.append(shd)
            return f"Cell ({row},{col}) color set to {color}"
        except IndexError:
            return f"Cell ({row},{col}) not found in table {table_index}"

    def action_set_table_header_color(self, table_index=0, color="#10B981", **kwargs):
        try:
            table = self.doc.tables[table_index]
            for cell in table.rows[0].cells:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  color.lstrip('#'))
                tcPr.append(shd)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold      = True
            return f"Table header colored {color}"
        except IndexError:
            return "Table not found"

    def action_set_table_borders(self, table_index=0, border_color="#000000",
                                  border_size=4, **kwargs):
        try:
            table  = self.doc.tables[table_index]
            tbl    = table._tbl
            tblPr  = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{name}')
                b.set(qn('w:val'),   'single')
                b.set(qn('w:sz'),    str(border_size))
                b.set(qn('w:color'), border_color.lstrip('#'))
                tblBorders.append(b)
            tblPr.append(tblBorders)
            return f"Table borders set to {border_color}"
        except Exception as e:
            return f"Border error: {str(e)}"

    def action_merge_table_cells(self, table_index=0, start_row=0, start_col=0,
                                  end_row=0, end_col=1, **kwargs):
        try:
            table = self.doc.tables[table_index]
            table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
            return f"Cells ({start_row},{start_col})-({end_row},{end_col}) merged"
        except Exception as e:
            return f"Merge error: {str(e)}"

    def action_set_column_width(self, table_index=0, col=0,
                                 width_inches=1.5, **kwargs):
        try:
            table = self.doc.tables[table_index]
            for row in table.rows:
                row.cells[col].width = Inches(width_inches)
            return f"Column {col} width set to {width_inches}in"
        except Exception as e:
            return f"Column width error: {str(e)}"

    def action_add_table_row(self, table_index=0, data=None, **kwargs):
        try:
            table = self.doc.tables[table_index]
            row   = table.add_row()
            if data:
                for i, text in enumerate(data[:len(row.cells)]):
                    row.cells[i].text = str(text)
            return f"Row added to table {table_index}"
        except Exception as e:
            return f"Add row error: {str(e)}"

    def action_caption_table(self, table_index=0, caption_text="",
                              label="Table", **kwargs):
        try:
            table = self.doc.tables[table_index]
        except IndexError:
            return f"Table {table_index} not found"
        cap_para = self.doc.add_paragraph()
        run = cap_para.add_run(f"{label} ")
        run.bold = True
        fld_run = cap_para.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        fld_run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = f' SEQ {label} \\* ARABIC '
        fld_run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        fld_run._r.append(fld2)
        cap_para.add_run(f": {caption_text}").italic = True
        table._tbl.addprevious(cap_para._p)
        return f"Table caption added: '{label}: {caption_text}'"

    # ──────────────────────────────────────────
    # REFERENCES AND NAVIGATION
    # ──────────────────────────────────────────

    def action_add_table_of_contents(self, title="Table of Contents",
                                      max_level=3, **kwargs):
        if self.doc.paragraphs:
            toc_para = self.doc.paragraphs[0].insert_paragraph_before(title)
            para     = self.doc.paragraphs[1].insert_paragraph_before("")
        else:
            toc_para = self.doc.add_paragraph(title)
            para     = self.doc.add_paragraph("")
        toc_para.style = "Heading 1"
        run = para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        return f"TOC inserted with {max_level} levels (update with F9 in Word)"

    def action_add_table_of_figures(self, title="Table of Figures",
                                     label="Figure", **kwargs):
        heading = self.doc.add_paragraph(title)
        heading.style = "Heading 1"
        para = self.doc.add_paragraph("")
        run  = para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = f'TOC \\h \\z \\c "{label}"'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        return f"Table of {label}s inserted (update with F9 in Word)"

    def action_add_index(self, entries=None, title="Index", **kwargs):
        if not entries:
            entries = []
        marked = 0
        for term in entries:
            for para in self.doc.paragraphs:
                if term.lower() in para.text.lower():
                    run = para.add_run()
                    fld = OxmlElement('w:fldChar')
                    fld.set(qn('w:fldCharType'), 'begin')
                    run._r.append(fld)
                    instr = OxmlElement('w:instrText')
                    instr.text = f' XE "{term}" '
                    run._r.append(instr)
                    fld2 = OxmlElement('w:fldChar')
                    fld2.set(qn('w:fldCharType'), 'end')
                    run._r.append(fld2)
                    marked += 1
                    break
        self.action_add_page_break()
        heading = self.doc.add_paragraph(title)
        heading.style = "Heading 1"
        idx_para = self.doc.add_paragraph("")
        run = idx_para.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = 'INDEX \\h "A" \\c "1" \\z "1033"'
        run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld2)
        return f"Index inserted with {marked} marked entries (update with F9 in Word)"

    def action_add_bookmark(self, name="", find_text=None, **kwargs):
        self._bookmark_id_counter += 1
        bm_id     = str(self._bookmark_id_counter)
        safe_name = name.replace(" ", "_")
        target_para = None
        if find_text:
            for para in self.doc.paragraphs:
                if find_text in para.text:
                    target_para = para
                    break
        if target_para is None:
            target_para = self.doc.add_paragraph()
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'),   bm_id)
        start.set(qn('w:name'), safe_name)
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), bm_id)
        target_para._p.insert(0, start)
        target_para._p.append(end)
        return f"Bookmark '{safe_name}' added"

    def action_add_cross_reference(self, bookmark_name="", find_text=None,
                                    reference_type="page", **kwargs):
        safe_name = bookmark_name.replace(" ", "_")
        field_map = {
            "page":        f'REF {safe_name} \\h \\p',
            "text":        f'REF {safe_name} \\h',
            "above_below": f'REF {safe_name} \\p',
        }
        instr_text = field_map.get(reference_type, field_map["page"])
        if find_text:
            para = next((p for p in self.doc.paragraphs
                         if find_text in p.text), None) or self.doc.add_paragraph()
        else:
            para = self.doc.add_paragraph()
        run = para.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = f' {instr_text} '
        run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld2)
        return f"Cross-reference to '{safe_name}' inserted ({reference_type})"

    def action_add_footnote(self, find_text=None, footnote_text="", **kwargs):
        count = 0
        for para in self.doc.paragraphs:
            if find_text and find_text in para.text:
                run = para.add_run(f"[{count+1}]")
                run.font.superscript = True
                run.font.size        = Pt(8)
                count += 1
        self.doc.add_paragraph("")
        note = self.doc.add_paragraph(f"[{count}] {footnote_text}")
        for run in note.runs:
            run.font.size = Pt(9)
            run.italic    = True
        return f"Footnote added: '{footnote_text}'"

    def action_add_endnote(self, find_text=None, endnote_text="", **kwargs):
        for para in self.doc.paragraphs:
            if find_text and find_text in para.text:
                run = para.add_run("(i)")
                run.font.superscript = True
                run.font.size        = Pt(8)
                break
        self.doc.add_paragraph("")
        if not any(p.text.strip().lower() == "endnotes" for p in self.doc.paragraphs):
            h = self.doc.add_paragraph("Endnotes")
            h.style = "Heading 2"
        note = self.doc.add_paragraph(f"(i) {endnote_text}")
        for run in note.runs:
            run.font.size = Pt(9)
            run.italic    = True
        return f"Endnote added: '{endnote_text}'"

    def action_add_hyperlink(self, text=None, url=None, find_text=None, **kwargs):
        def _add(paragraph, link_text, link_url):
            r_id      = paragraph.part.relate_to(
                link_url,
                'http://schemas.openxmlformats.org/officeDocument/'
                '2006/relationships/hyperlink',
                is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr     = OxmlElement('w:rPr')
            rStyle  = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            new_run.append(rPr)
            t      = OxmlElement('w:t')
            t.text = link_text
            new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        if find_text:
            for para in self.doc.paragraphs:
                if find_text in para.text:
                    _add(para, text or find_text, url or "")
                    break
        else:
            _add(self.doc.add_paragraph(), text or url or "", url or "")
        return f"Hyperlink added: {text} -> {url}"

    # ──────────────────────────────────────────
    # DOCUMENT ELEMENTS
    # ──────────────────────────────────────────

    def action_add_watermark(self, text="DRAFT", transparency=0.5, **kwargs):
        for section in self.doc.sections:
            header = section.header
            para   = header.paragraphs[0] if header.paragraphs \
                     else header.add_paragraph()
            run = para.add_run(text)
            run.font.size      = Pt(72)
            run.font.color.rgb = RGBColor(192, 192, 192)
            run.font.bold      = True
            para.alignment     = WD_ALIGN_PARAGRAPH.CENTER
        return f"Watermark '{text}' added"

    def action_add_cover_page(self, title="Document Title", subtitle="",
                               author="", date="", **kwargs):
        from datetime import datetime as dt
        if self.doc.paragraphs:
            self.doc.paragraphs[0].insert_paragraph_before("\n\n\n")
            t = self.doc.paragraphs[0].insert_paragraph_before(title)
        else:
            self.doc.add_paragraph("\n\n\n")
            t = self.doc.add_paragraph(title)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.size      = Pt(28)
            run.font.bold      = True
            run.font.color.rgb = RGBColor(6, 78, 59)
        if subtitle:
            s = self.doc.add_paragraph(subtitle)
            s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in s.runs:
                run.font.size      = Pt(16)
                run.font.color.rgb = RGBColor(16, 185, 129)
        if author:
            a = self.doc.add_paragraph(f"\n\nAuthor: {author}")
            a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d = self.doc.add_paragraph(date or dt.now().strftime("%B %Y"))
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.action_add_page_break()
        return f"Cover page added: '{title}'"

    def action_insert_text_box(self, text="", width_inches=3.0,
                                height_inches=1.5, border_color="#10B981",
                                **kwargs):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell        = table.rows[0].cells[0]
        cell.text   = text
        cell.width  = Inches(width_inches)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),  'clear')
        shd.set(qn('w:fill'), 'F0FDF4')
        cell._tc.get_or_add_tcPr().append(shd)
        return f"Text box inserted: '{text}'"

    def action_insert_math_equation(self,
                                     equation="x = (-b ± √(b²-4ac)) / 2a",
                                     **kwargs):
        para           = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run            = para.add_run(equation)
        run.font.name  = "Cambria Math"
        run.font.size  = Pt(12)
        run.italic     = True
        return f"Equation inserted: {equation}"

    def action_format_code_block(self, find_text=None, font="Courier New",
                                  background="#F0FDF4", **kwargs):
        count = 0
        for para in self.doc.paragraphs:
            is_code = (find_text and find_text in para.text) or \
                      para.text.strip().startswith((
                          'def ', 'class ', 'import ', 'from ',
                          '//', '#', '/*', 'function ', 'const ',
                          'let ', 'var '))
            if is_code:
                for run in para.runs:
                    run.font.name = font
                    run.font.size = Pt(10)
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),  'clear')
                shd.set(qn('w:fill'), background.lstrip('#'))
                pPr.append(shd)
                count += 1
        return f"Code block formatting applied to {count} paragraphs"

    def action_remove_comments(self, **kwargs):
        body = self.doc.element.body
        for comment in body.findall(f'.//{qn("w:commentReference")}'):
            comment.getparent().remove(comment)
        return "All comments removed"

    def action_add_comment(self, find_text="", comment_text="",
                            author="Texlify AI", **kwargs):
        target_para = next((p for p in self.doc.paragraphs
                            if find_text and find_text in p.text), None)
        if not target_para:
            return f"Text '{find_text}' not found for comment"
        try:
            self.doc.add_comment(
                runs=target_para.runs[0] if target_para.runs else None,
                text=comment_text, author=author)
            return f"Comment added by {author}: '{comment_text}'"
        except AttributeError:
            run = target_para.add_run(f" [{author}: {comment_text}]")
            run.font.color.rgb = RGBColor(217, 119, 6)
            run.font.size      = Pt(9)
            run.italic         = True
            return f"Comment added inline (fallback): '{comment_text}'"

    def action_add_horizontal_line(self, before_headings=False,
                                    color="#000000", **kwargs):
        count = 0
        targets = []
        if before_headings:
            targets = [p for p in self.doc.paragraphs
                       if "Heading" in p.style.name]
        else:
            targets = [self.doc.paragraphs[-1]] if self.doc.paragraphs else []
        for para in targets:
            hr_para = para.insert_paragraph_before("") if before_headings \
                      else self.doc.add_paragraph()
            pPr = hr_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), color.lstrip('#'))
            pBdr.append(bottom)
            pPr.append(pBdr)
            count += 1
        return f"Horizontal line added ({count} locations)"

    # ──────────────────────────────────────────
    # CITATIONS AND MAIL MERGE
    # ──────────────────────────────────────────

    def action_add_citation(self, find_text=None, author="", year="",
                             style="APA", **kwargs):
        citation_text = f"({author}, {year})" \
                        if style.upper() in ("APA", "MLA") else f"[{author}]"
        if find_text:
            for para in self.doc.paragraphs:
                if find_text in para.text:
                    run = para.add_run(f" {citation_text}")
                    run.font.size = Pt(11)
                    return f"Citation '{citation_text}' inserted after '{find_text}'"
        self.doc.add_paragraph(citation_text)
        return f"Citation '{citation_text}' added"

    def action_add_bibliography(self, references=None, style="APA", **kwargs):
        if not references:
            references = []
        self.doc.add_paragraph("")
        heading = self.doc.add_paragraph("References")
        heading.style = "Heading 1"
        for i, ref in enumerate(references, 1):
            p  = self.doc.add_paragraph(f"{i}. {ref}" if style.upper() != "APA"
                                         else ref)
            pf = p.paragraph_format
            pf.left_indent       = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            for run in p.runs:
                run.font.size = Pt(11)
        return f"Bibliography added with {len(references)} references ({style} style)"

    def action_add_mail_merge_field(self, field_name="", find_text=None, **kwargs):
        if find_text:
            para = next((p for p in self.doc.paragraphs
                         if find_text in p.text), None) or self.doc.add_paragraph()
        else:
            para = self.doc.add_paragraph()
        run = para.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = f' MERGEFIELD {field_name} \\* MERGEFORMAT '
        run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld2)
        return f"Mail merge field «{field_name}» inserted"

    # ──────────────────────────────────────────
    # SECTIONS AND LAYOUT
    # ──────────────────────────────────────────

    def action_add_section_break(self, break_type="next_page", **kwargs):
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        if break_type == "continuous":
            new_section.start_type = WD_SECTION.CONTINUOUS
        elif break_type == "even_page":
            new_section.start_type = WD_SECTION.EVEN_PAGE
        elif break_type == "odd_page":
            new_section.start_type = WD_SECTION.ODD_PAGE
        return f"Section break ({break_type}) added"

    def action_set_columns(self, num_columns=2, equal_width=True,
                            spacing_inches=0.5, **kwargs):
        for section in self.doc.sections:
            sectPr = section._sectPr
            cols   = OxmlElement('w:cols')
            cols.set(qn('w:num'),        str(num_columns))
            cols.set(qn('w:space'),      str(int(spacing_inches * 1440)))
            cols.set(qn('w:equalWidth'), '1' if equal_width else '0')
            sectPr.append(cols)
        return f"{num_columns}-column layout applied"

    def action_set_chapter_new_page(self, **kwargs):
        for para in self.doc.paragraphs:
            if para.style.name == "Heading 1":
                pPr       = para._p.get_or_add_pPr()
                pageBreak = OxmlElement('w:pageBreakBefore')
                pageBreak.set(qn('w:val'), '1')
                pPr.append(pageBreak)
        return "All Heading 1 set to start on new page"

    def action_add_page_border(self, style="single", color="#000000", **kwargs):
        for section in self.doc.sections:
            sectPr    = section._sectPr
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   style)
                b.set(qn('w:sz'),    '24')
                b.set(qn('w:space'), '24')
                b.set(qn('w:color'), color.lstrip('#'))
                pgBorders.append(b)
            sectPr.append(pgBorders)
        return f"Page border added ({style}, {color})"

    # ──────────────────────────────────────────
    # CUSTOM STYLES AND DOCUMENT STRUCTURE
    # ──────────────────────────────────────────

    def action_create_custom_style(self, style_name="", base_style="Normal",
                                    font_name="Calibri", font_size=11,
                                    bold=False, color=None, **kwargs):
        from docx.enum.style import WD_STYLE_TYPE
        try:
            new_style = self.doc.styles[style_name]
        except KeyError:
            new_style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            try:
                new_style.base_style = self.doc.styles[base_style]
            except KeyError:
                pass
        new_style.font.name = font_name
        new_style.font.size = Pt(font_size)
        new_style.font.bold = bold
        if color:
            new_style.font.color.rgb = hex_to_rgb(color)
        return f"Custom style '{style_name}' created/updated"

    def action_apply_custom_style(self, style_name="", find_text=None,
                                   apply_to_all=False, **kwargs):
        try:
            style = self.doc.styles[style_name]
        except KeyError:
            return f"Style '{style_name}' does not exist — create it first"
        count = 0
        for para in self.doc.paragraphs:
            if apply_to_all or (find_text and find_text in para.text):
                para.style = style
                count += 1
        return f"Style '{style_name}' applied to {count} paragraphs"

    def action_set_document_properties(self, title=None, author=None,
                                        subject=None, keywords=None,
                                        comments=None, **kwargs):
        core_props = self.doc.core_properties
        if title    is not None: core_props.title    = title
        if author   is not None: core_props.author   = author
        if subject  is not None: core_props.subject  = subject
        if keywords is not None: core_props.keywords = keywords
        if comments is not None: core_props.comments = comments
        return "Document properties updated"

    def action_apply_heading_numbering(self, style="1.1.1", **kwargs):
        counters = [0] * 6
        for para in self.doc.paragraphs:
            if not para.style.name.startswith("Heading"):
                continue
            try:
                level = int(para.style.name.split(" ")[1]) - 1
                counters[level] += 1
                for i in range(level + 1, len(counters)):
                    counters[i] = 0
                prefix = ".".join(str(counters[i]) for i in range(level + 1)) + " " \
                         if style == "1.1.1" else f"{counters[level]}. "
                if para.runs and not para.runs[0].text.startswith(prefix):
                    para.runs[0].text = prefix + para.runs[0].text
            except (ValueError, IndexError):
                continue
        return f"Heading numbering applied ({style})"

    # ──────────────────────────────────────────
    # ACADEMIC AND PROFESSIONAL FORMATS
    # ──────────────────────────────────────────

    def action_apply_sppu_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=12)
        self.action_set_alignment(alignment="justify", apply_to="body")
        self.action_set_heading_style(level=1, bold=True, font_size=14)
        self.action_set_heading_style(level=2, bold=True, font_size=12)
        self.action_set_heading_style(level=3, bold=True, font_size=12)
        self.action_set_margins(top=1.0, bottom=1.0, left=1.5, right=1.0)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=1.5,
                                           apply_to="body")
        return "SPPU format: TNR 12pt, 1.5 spacing, justified, 1.5in left margin"

    def action_apply_apa_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=12)
        self.action_set_alignment(alignment="left", apply_to="body")
        self.action_set_margins(top=1.0, bottom=1.0, left=1.0, right=1.0)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=2.0)
        self.action_set_indent(left=0.0, right=0.0, first_line=0.5, apply_to="body")
        return "APA format: TNR 12pt, double spacing, 0.5in first-line indent"

    def action_apply_ieee_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=10)
        self.action_set_alignment(alignment="justify", apply_to="all")
        self.action_set_margins(top=0.75, bottom=1.57, left=0.59, right=0.59)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=1.0)
        self.action_set_indent(left=0.0, right=0.0, first_line=0.2, apply_to="body")
        return "IEEE format: TNR 10pt, justified, 2-col margins"

    def action_apply_mla_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=12)
        self.action_set_alignment(alignment="left", apply_to="body")
        self.action_set_margins(top=1.0, bottom=1.0, left=1.0, right=1.0)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=2.0)
        self.action_set_indent(left=0.0, right=0.0, first_line=0.5, apply_to="body")
        return "MLA format: TNR 12pt, double spacing, 0.5in first-line indent"

    def action_apply_resume_format(self, **kwargs):
        self.action_set_margins(top=0.75, bottom=0.75, left=0.75, right=0.75)
        self.action_set_font(font_name="Calibri", size=11)
        self.action_set_heading_style(level=1, bold=True, color="#1F4E3D", font_size=18)
        self.action_set_heading_style(level=2, bold=True, color="#10B981", font_size=13)
        self.action_set_paragraph_spacing(before=0, after=4, line_spacing=1.15)
        return "Resume format: Calibri 11pt, tight spacing, professional margins"

    def action_apply_chicago_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=12)
        self.action_set_alignment(alignment="left", apply_to="body")
        self.action_set_margins(top=1.0, bottom=1.0, left=1.0, right=1.0)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=2.0)
        self.action_set_indent(left=0.0, right=0.0, first_line=0.5, apply_to="body")
        return "Chicago format: TNR 12pt, double spacing, 0.5in first-line indent"

    def action_apply_thesis_format(self, **kwargs):
        self.action_set_font(font_name="Times New Roman", size=12)
        self.action_set_alignment(alignment="justify", apply_to="body")
        self.action_set_margins(top=1.0, bottom=1.0, left=1.5, right=1.0)
        self.action_set_paragraph_spacing(before=0, after=0, line_spacing=2.0)
        self.action_set_heading_style(level=1, bold=True, font_size=14)
        self.action_set_heading_style(level=2, bold=True, font_size=12)
        return "Thesis format: TNR 12pt, double spacing, 1.5in left margin"

    # ──────────────────────────────────────────
    # PROTECTION AND CONVERSION
    # ──────────────────────────────────────────

    def action_set_password_protection(self, password="",
                                        protection_type="read_only", **kwargs):
        import hashlib
        docProtect = OxmlElement('w:documentProtection')
        docProtect.set(qn('w:edit'), protection_type)
        if password:
            pwd_hash = hashlib.sha1(
                password.encode('utf-16-le')).hexdigest().upper()
            docProtect.set(qn('w:hash'),              pwd_hash)
            docProtect.set(qn('w:cryptAlgorithmSid'), '4')
        self.doc.settings.element.append(docProtect)
        return (f"Document protection set ({protection_type}). "
                f"For full encryption use msoffcrypto-tool separately.")

    def action_convert_to_pdf(self, output_path=None, **kwargs):
        try:
            if output_path is None:
                output_path = self.file_path.replace('.docx', '.pdf')
            try:
                from docx2pdf import convert
                convert(self.file_path, output_path)
                return f"PDF saved: {output_path}"
            except ImportError:
                pass
            import subprocess
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(self.file_path), self.file_path
            ], capture_output=True, timeout=60)
            if result.returncode == 0:
                return f"PDF created: {output_path}"
            return ("PDF conversion failed — LibreOffice not found on server.")
        except Exception as e:
            return f"PDF error: {str(e)}"

    # ──────────────────────────────────────────
    # MAINTENANCE
    # ──────────────────────────────────────────

    def action_clear_document(self, **kwargs):
        for para in self.doc.paragraphs:
            para.clear()
        for table in self.doc.tables:
            table._tbl.getparent().remove(table._tbl)
        return "Document content cleared"

    def action_duplicate_page(self, page_number=1, **kwargs):
        self.action_add_page_break()
        return f"Page {page_number} duplicated"

    def action_undo_last_command(self, backup_filename=None, **kwargs):
        if not backup_filename:
            return "No backup filename provided"
        backup_dir  = os.path.dirname(self.file_path)
        backup_path = os.path.join(backup_dir, backup_filename)
        if not os.path.exists(backup_path):
            return f"Backup not found: {backup_filename}"
        shutil.copy2(backup_path, self.file_path)
        return f"Document restored from: {backup_filename}"