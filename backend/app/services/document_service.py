import os
import uuid
import base64
import glob
import logging
import mammoth
from fastapi import HTTPException, UploadFile, status
from sqlalchemy.orm import Session
from app.models.document import Document
from app.models.user import User
from app.core.config import settings
from docx import Document as DocxDocument

logger = logging.getLogger("texlify.documents")

ALLOWED_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
]


def _convert_image(image):
    try:
        with image.open() as img_file:
            image_bytes = img_file.read()
        encoded = base64.b64encode(image_bytes).decode('utf-8')
        return {"src": f"data:{image.content_type};base64,{encoded}"}
    except Exception:
        return {"src": ""}


def _safe_path(path: str, base_dir: str) -> str:
    """Resolve path and assert it stays inside base_dir. Raises 403 if not."""
    resolved = os.path.realpath(path)
    base     = os.path.realpath(base_dir)
    if not resolved.startswith(base + os.sep) and resolved != base:
        raise HTTPException(status_code=403, detail="Access denied")
    return resolved


def _cleanup_old_backups(file_path: str, keep: int = None):
    """Keep only the N most recent backups for a document."""
    if keep is None:
        keep = settings.MAX_BACKUPS_PER_DOCUMENT
    pattern = f"{file_path}.backup_*"
    backups = sorted(glob.glob(pattern))
    for old in backups[:-keep] if keep > 0 else backups:
        try:
            os.remove(old)
        except OSError:
            pass


class DocumentService:

    @staticmethod
    def get_user_upload_dir(user_id: int) -> str:
        path = os.path.join(settings.UPLOAD_DIR, str(user_id))
        os.makedirs(path, exist_ok=True)
        return path

    @staticmethod
    def get_page_count(file_path: str) -> int:
        try:
            doc = DocxDocument(file_path)
            return max(1, len(doc.paragraphs) // 25)
        except Exception:
            return 1

    @staticmethod
    def convert_to_html(file_path: str) -> str:
        try:
            style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Heading 4'] => h4:fresh
                p[style-name='Heading 5'] => h5:fresh
                p[style-name='Heading 6'] => h6:fresh
                p[style-name='List Paragraph'] => li:fresh
                p[style-name='List Bullet']    => li.bullet:fresh
                p[style-name='List Number']    => li.numbered:fresh
                u => u
                b => strong
                i => em
                strike => s
                del => s
            """
            with open(file_path, "rb") as f:
                result = mammoth.convert_to_html(
                    f,
                    style_map=style_map,
                    convert_image=mammoth.images.img_element(_convert_image)
                )
            html_body = result.value
            return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: 'Times New Roman', Times, serif;
    font-size: 12pt; line-height: 1.6;
    color: #000; background: #fff;
    padding: 40px 60px; max-width: 800px; margin: 0 auto;
  }}
  h1 {{ font-size: 18pt; font-weight: bold; margin: 16px 0 8px; color: #1a1a1a; }}
  h2 {{ font-size: 15pt; font-weight: bold; margin: 14px 0 6px; color: #1a1a1a; }}
  h3 {{ font-size: 13pt; font-weight: bold; margin: 12px 0 4px; color: #1a1a1a; }}
  h4, h5, h6 {{ font-size: 12pt; font-weight: bold; margin: 10px 0 4px; }}
  p  {{ margin: 6px 0; text-align: justify; }}
  u  {{ text-decoration: underline; }}
  strong, b {{ font-weight: bold; }}
  em, i {{ font-style: italic; }}
  s, strike, del {{ text-decoration: line-through; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  td, th {{ border: 1px solid #000; padding: 6px 10px; font-size: 11pt; }}
  th {{ background: #f0f0f0; font-weight: bold; }}
  ul, ol {{ margin: 8px 0 8px 24px; }}
  li {{ margin: 3px 0; }}
  img {{ max-width: 100%; height: auto; margin: 8px 0; display: block; }}
  hr {{ border: none; border-top: 1px solid #000; margin: 12px 0; }}
  a  {{ color: #10B981; text-decoration: underline; }}
  sup {{ font-size: 75%; vertical-align: super; }}
  sub {{ font-size: 75%; vertical-align: sub; }}
  pre, code {{ font-family: 'Courier New', monospace; font-size: 10pt;
               background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
  pre {{ padding: 12px; margin: 8px 0; overflow-x: auto; }}
  blockquote {{ border-left: 3px solid #10B981; padding-left: 16px;
                margin: 8px 0 8px 16px; color: #555; font-style: italic; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
        except Exception as e:
            logger.error("Preview conversion failed for %s: %s", file_path, e)
            return f"""<!DOCTYPE html>
<html><body style="font-family:sans-serif;padding:40px;color:#666;">
  <p>Preview unavailable: {str(e)}</p>
  <p>Download the document to see the full formatting.</p>
</body></html>"""

    @staticmethod
    def upload(db: Session, user: User, file: UploadFile) -> Document:
        if file.content_type not in ALLOWED_TYPES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only .docx or .doc files are allowed"
            )
        contents  = file.file.read()
        file_size = len(contents)
        if file_size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File size exceeds {settings.MAX_FILE_SIZE_MB}MB limit"
            )
        ext             = os.path.splitext(file.filename)[1]
        stored_filename = f"{uuid.uuid4().hex}{ext}"
        upload_dir      = DocumentService.get_user_upload_dir(user.id)
        file_path       = os.path.join(upload_dir, stored_filename)

        # Path traversal check
        _safe_path(file_path, upload_dir)

        with open(file_path, "wb") as f:
            f.write(contents)
        title = (
            os.path.splitext(file.filename)[0]
            .replace("_", " ").replace("-", " ").title()
        )
        document = Document(
            user_id=user.id,
            title=title,
            original_filename=file.filename,
            stored_filename=stored_filename,
            file_path=file_path,
            file_size=file_size,
            page_count=DocumentService.get_page_count(file_path),
        )
        db.add(document); db.commit(); db.refresh(document)
        logger.info("Document uploaded: user=%s file=%s size=%d",
                    user.id, stored_filename, file_size)
        return document

    @staticmethod
    def get_all(db: Session, user: User) -> list:
        return (
            db.query(Document)
            .filter(Document.user_id == user.id, Document.is_active == True)
            .order_by(Document.created_at.desc())
            .all()
        )

    @staticmethod
    def get_by_id(db: Session, document_id: int, user: User) -> Document:
        document = (
            db.query(Document)
            .filter(
                Document.id        == document_id,
                Document.user_id   == user.id,
                Document.is_active == True
            )
            .first()
        )
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        return document

    @staticmethod
    def delete(db: Session, document_id: int, user: User) -> bool:
        document = DocumentService.get_by_id(db, document_id, user)
        upload_dir = DocumentService.get_user_upload_dir(user.id)
        if os.path.exists(document.file_path):
            _safe_path(document.file_path, upload_dir)
            # Clean up all backups too
            _cleanup_old_backups(document.file_path, keep=0)
            os.remove(document.file_path)
        document.is_active = False
        db.commit()
        return True